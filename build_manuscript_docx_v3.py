"""Build the restructured manuscript DOCX: main text, then Figure Legends,
then all Figures, then all Tables. Figure placeholders are *[file.png]*."""
import glob
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

SRC = "/home/ubuntu/implant_manuscript_v3.md"
OUT = "/home/ubuntu/implant_manuscript_v3.docx"

FIGS = {os.path.basename(p): p for p in glob.glob("/home/ubuntu/*.png")}

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def add_runs(par, text):
    for part in re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            par.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            par.add_run(part[1:-1]).font.name = "Consolas"
        else:
            par.add_run(part)


with open(SRC) as fh:
    lines = fh.read().split("\n")

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    stripped = line.strip()

    if stripped == "<<<PAGEBREAK>>>":
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        i += 1
        continue

    fig = re.search(r"\[([\w\.\-]+\.png)\]", line)
    if fig and fig.group(1) in FIGS:
        label = re.sub(r"\*?\[[\w\.\-]+\.png\]\*?", "", line).replace("**", "").strip()
        if label:
            par = doc.add_paragraph()
            par.add_run(label).bold = True
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(FIGS[fig.group(1)], width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        i += 1
        continue

    if stripped.startswith("|") and i + 1 < len(lines) and \
            re.match(r"^\s*\|[\s:\-\|]+\|\s*$", lines[i + 1]):
        header = [c.strip() for c in stripped.strip("|").split("|")]
        rows = []
        j = i + 2
        while j < len(lines) and lines[j].strip().startswith("|"):
            rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
            j += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Light Grid Accent 1"
        for k, head in enumerate(header):
            cell = table.rows[0].cells[k]
            cell.paragraphs[0].clear()
            add_runs(cell.paragraphs[0], head)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
        for row in rows:
            cells = table.add_row().cells
            for k, val in enumerate(row):
                if k < len(cells):
                    cells[k].paragraphs[0].clear()
                    add_runs(cells[k].paragraphs[0], val)
                    for run in cells[k].paragraphs[0].runs:
                        run.font.size = Pt(8)
        doc.add_paragraph()
        i = j
        continue

    if stripped == "---":
        i += 1
        continue

    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=0)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
    elif line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
    elif stripped:
        add_runs(doc.add_paragraph(), line)
    i += 1

doc.save(OUT)
print("wrote", OUT)
